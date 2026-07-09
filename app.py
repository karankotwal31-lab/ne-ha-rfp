import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from pypdf import PdfReader
import os
import re
import json
import sqlite3
import pandas as pd
from docx import Document
from openpyxl import Workbook
from io import BytesIO

# Set global structural app page configurations
st.set_page_config(
    page_title="Ne-Ha | Enterprise RFP Data Matrix", 
    page_icon="🚀", 
    layout="wide", 
    initial_sidebar_state="expanded"
)

DB_FILE = "neha_enterprise_rfp.db"

# --- ENTERPRISE RELATIONAL RETENTION LAYER ---
def init_db():
    """Establishes localized tables to protect data from browser tab timeouts."""
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS rfp_matrix (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            doc_name TEXT,
            section TEXT,
            requirement_extracted TEXT,
            category TEXT,
            assigned_team TEXT,
            status TEXT,
            risk_multiplier REAL,
            historical_match TEXT
        )
    """)
    conn.commit()
    conn.close()

def save_dataframe_to_db(df, doc_name):
    """Safely writes the current dataframe matrix down to relational database storage."""
    if df is None or df.empty:
        return
    conn = sqlite3.connect(DB_FILE)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM rfp_matrix WHERE doc_name = ?", (doc_name,))
    
    df_to_save = df.copy()
    df_to_save["doc_name"] = doc_name
    df_to_save.to_sql("rfp_matrix", conn, if_exists="append", index=False)
    conn.commit()
    conn.close()

def load_latest_document_from_db(doc_name):
    """Restores active screen metrics natively if a user session breaks or sleeps."""
    conn = sqlite3.connect(DB_FILE)
    query = "SELECT section, requirement_extracted, category, assigned_team, status, risk_multiplier, historical_match FROM rfp_matrix WHERE doc_name = ?"
    df = pd.read_sql_query(query, conn, params=(doc_name,))
    conn.close()
    return df if not df.empty else None

# --- PHASE 3.5: LAYOUT TEXT SANITIZER ---
def robust_text_sanitizer(text):
    if not text:
        return ""
    text = re.sub(r'(?<=\b\w)\s(?=\w\b)', '', text)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)
    text = re.sub(r'\n\s*\n', '\n', text)
    return text.strip()

# --- HIGH-RESILIENCE JSON STRUCTURAL REPAIR ENGINE ---
def extract_json_array(response_text):
    if not response_text:
        return get_emergency_fallback("Null string payload passed down to parser.")
        
    cleaned_text = response_text.strip()
    cleaned_text = re.sub(r'^```json\s*', '', cleaned_text, flags=re.IGNORECASE)
    cleaned_text = re.sub(r'^```\s*', '', cleaned_text, flags=re.IGNORECASE)
    cleaned_text = re.sub(r'\s*```$', '', cleaned_text, flags=re.IGNORECASE)
    cleaned_text = cleaned_text.strip()

    try:
        match = re.search(r'\[\s*\{.*\}\s*\]', cleaned_text, re.DOTALL)
        if match:
            return json.loads(match.group(0))
        return json.loads(cleaned_text)
    except Exception as e:
        try:
            fixed_text = re.sub(r'(?<=:\s")(.*?)(?=",?\s*\n)', lambda m: m.group(1).replace('"', '\\"'), cleaned_text)
            match = re.search(r'\[\s*\{.*\}\s*\]', fixed_text, re.DOTALL)
            if match:
                return json.loads(match.group(0))
        except Exception:
            pass
        return get_emergency_fallback(f"JSON Framework Exception: {str(e)}")

def get_emergency_fallback(reason):
    return [{
        "section": "Raw Extraction Fallback",
        "requirement_extracted": f"Parser Security Intercept: {reason}",
        "category": "Proposal Draft",
        "assigned_team": "Sales Planning",
        "status": "Pending Review",
        "risk_multiplier": 1.00,
        "historical_match": "0% (State Recovery Core Active)"
    }]

# --- ASYNCHRONOUS COGNITIVE CELL MANAGER ---
def sync_matrix_edits():
    """Callback function that cleanly updates the database when a cell is edited."""
    if "interactive_matrix_editor" in st.session_state and st.session_state.get("active_doc_name"):
        edits = st.session_state["interactive_matrix_editor"]
        df = st.session_state["rfp_data_matrix"]
        doc_name = st.session_state["active_doc_name"]
        
        # Process modifications smoothly
        for row_idx, changed_cols in edits.get("edited_rows", {}).items():
            for col_name, new_val in changed_cols.items():
                df.iat[row_idx, df.columns.get_loc(col_name)] = new_val
                
        # Handle new manual row assignments inside UI matrix
        for added_row in edits.get("added_rows", {}):
            df = pd.concat([df, pd.DataFrame([added_row])], ignore_index=True)
            
        st.session_state["rfp_data_matrix"] = df
        save_dataframe_to_db(df, doc_name)

# --- BI-DIRECTIONAL EXCEL LAYOUT MATRIX STYLER ---
def generate_enterprise_excel(df):
    """Converts local active runtime updates into an enterprise-ready formatted Excel document."""
    output = BytesIO()
    wb = Workbook()
    ws = wb.active
    ws.title = "RFP Requirements Matrix"
    ws.views.sheetView[0].showGridLines = True
    
    headers = list(df.columns)
    ws.append(headers)
    
    for _, row in df.iterrows():
        ws.append(list(row))
        
    wb.save(output)
    return output.getvalue()

# Initialize background structural dependencies
init_db()

if "rfp_data_matrix" not in st.session_state:
    st.session_state["rfp_data_matrix"] = None
if "active_doc_name" not in st.session_state:
    st.session_state["active_doc_name"] = None

# Sidebar Console Layout Control Panel
with st.sidebar:
    st.title("⚙️ Commercial Panel")
    st.markdown("Welcome to **Ne-Ha**, your autonomous workspace.")
    st.write("---")
    api_key = st.text_input("Gemini API Key:", type="password", help="Enter your Google AI developer key.")
    st.write("---")
    st.markdown("### Workspace Telemetry")
    if api_key:
        st.success("Core Engine Online")
    else:
        st.warning("Awaiting Authorization")

st.title("🚀 Ne-Ha Portal")
st.caption("Universal Multi-Format Commercial RFP Matrix & Workflow Workspace [Stateful Layer v5.5]")
st.write("---")

col_left, col_right = st.columns([1, 1.5], gap="large")

with col_left:
    st.subheader("📁 Snap & Audit: Document Ingestion")
    st.markdown("Drop client procurement spreadsheets (.xlsx), narrative contracts (.docx), or legacy PDFs below.")
    
    # UPGRADE: Comprehensive multi-format ingestion accept flags
    uploaded_file = st.file_uploader(
        "Upload Client Commercial File:", 
        type=["pdf", "docx", "xlsx"], 
        key="rfp_uploader"
    )
    
    if uploaded_file is not None:
        st.info(f"📄 File Target Ingested: '{uploaded_file.name}'")
        if st.session_state["active_doc_name"] != uploaded_file.name:
            cached_df = load_latest_document_from_db(uploaded_file.name)
            if cached_df is not None:
                st.session_state["rfp_data_matrix"] = cached_df
                st.session_state["active_doc_name"] = uploaded_file.name
                st.toast("Auto-Restored Workspace State from Local DB Archive! 🎉")

    generate_btn = st.button("Execute Data Matrix Shredder", type="primary", use_container_width=True)

with col_right:
    st.subheader("✨ Enterprise Requirement Workspace Matrix")
    
    if generate_btn:
        if not api_key:
            st.error("🔑 Configuration missing: Please input your Gemini API Key in the left sidebar.")
        elif uploaded_file is None:
            st.error("❌ Document missing: Please upload a file on the left before running execution.")
        else:
            try:
                os.environ["GOOGLE_API_KEY"] = api_key
                st.session_state["active_doc_name"] = uploaded_file.name
                
                llm = ChatGoogleGenerativeAI(
                    model="gemini-2.5-flash", 
                    transport="rest",
                    temperature=0.1
                )
                
                raw_extracted_text = ""
                file_extension = uploaded_file.name.split(".")[-1].lower()
                
                with st.spinner(f"⏳ Processing layout layers from .{file_extension} binary asset..."):
                    
                    # Core Handler A: Advanced PDF String Extractor
                    if file_extension == "pdf":
                        reader = PdfReader(uploaded_file)
                        target_pages = reader.pages[:40] 
                        for page in target_pages:
                            try:
                                page_text = page.extract_text()
                                if page_text:
                                    raw_extracted_text += page_text + "\n"
                            except Exception:
                                continue
                                
                    # Core Handler B: Document Table Merging-Safe Parser
                    elif file_extension == "docx":
                        doc = Document(uploaded_file)
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                raw_extracted_text += paragraph.text + "\n"
                        
                        # Defensively extract tables to eliminate duplicate strings from merged cells
                        for table in doc.tables:
                            seen_cells = set()
                            for row in table.rows:
                                row_text = []
                                for cell in row.cells:
                                    if cell not in seen_cells:
                                        seen_cells.add(cell)
                                        txt = cell.text.strip()
                                        if txt:
                                            row_text.append(txt)
                                if row_text:
                                    raw_extracted_text += " | ".join(row_text) + "\n"
                                    
                    # Core Handler C: High-Fidelity Excel to Markdown Matrix Transpiler
                    elif file_extension == "xlsx":
                        # Read all sheets to ensure hidden spreadsheet matrices aren't skipped
                        excel_file = pd.ExcelFile(uploaded_file)
                        for sheet_name in excel_file.sheet_names:
                            excel_df = pd.read_excel(uploaded_file, sheet_name=sheet_name)
                            if not excel_df.empty:
                                raw_extracted_text += f"\n--- Sheet: {sheet_name} ---\n"
                                # Convert rows into clean Markdown tables for reliable context window parsing
                                raw_extracted_text += excel_df.to_markdown(index=False) + "\n"
                
                cleaned_rfp_text = robust_text_sanitizer(raw_extracted_text)
                
                if not cleaned_rfp_text:
                    st.error("⚠️ Ingestion Failure: This document contains no valid readable text coordinates.")
                else:
                    with st.spinner("🧠 Orchestrating Structural Requirement Extraction Matrix..."):
                        
                        prompt = f"""You are an elite corporate proposal specialist, risk auditor, and technical software engineer.
                        Analyze the text block below. Core objective: extract the critical, most meaningful 15-20 core execution line-items, commercial conditions, load limits, or liabilities.
                        
                        CRITICAL FORMATTING BOUNDARY: You must return ONLY a clean, valid raw JSON array block. Do NOT use markdown symbols, do NOT include backticks (```json), and escape any interior quotation symbols.
                        
                        Expected Schema Framework:
                        {{
                            "section": "Clause reference or section header identified in document",
                            "requirement_extracted": "Granular summary of requirement, timeline demand, or equipment expectation",
                            "category": "Must be exactly one of: 'Proposal Draft', 'Risk Audit', or 'Calculation Layer'",
                            "assigned_team": "Must be exactly one of: 'Sales Planning', 'Legal/Compliance', 'Technical Ops', or 'Commercial Finance'",
                            "status": "Pending Review",
                            "risk_multiplier": 1.00,
                            "historical_match": "92% Match to Historical Bid"
                        }}

                        --- EXTRACTED CLIENT RFP TEXT ---
                        {cleaned_rfp_text}
                        """
                        
                        try:
                            response_object = llm.invoke(prompt)
                            full_output = response_object.content
                        except Exception as api_err:
                            raise RuntimeError(f"Cognitive Inference Boundary Violation: {api_err}")
                        
                        parsed_json_data = extract_json_array(full_output)
                        compiled_df = pd.DataFrame(parsed_json_data)
                        
                        st.session_state["rfp_data_matrix"] = compiled_df
                        save_dataframe_to_db(compiled_df, uploaded_file.name)
                        st.toast("Stateful Data Grid Compiled & Archived!", icon="🚀")

            except Exception as e:
                st.error(f"Execution Error Caught & Isolated: {e}")

    # --- DECOUPLED RENDERING ENVIRONMENT ---
    if st.session_state["rfp_data_matrix"] is not None:
        current_df = st.session_state["rfp_data_matrix"]
        
        # Real-time Telemetry Dashboard calculations
        high_risk_rows = current_df[current_df["risk_multiplier"] > 1.00] if "risk_multiplier" in current_df.columns else []
        total_items = len(current_df)
        risk_percentage = round((len(high_risk_rows) / total_items) * 100) if total_items > 0 else 0
        
        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric("Total Extracted Lines", f"{total_items} Items")
        with c2:
            st.metric("Risk Incidents Flagged", f"{len(high_risk_rows)} Rows")
        with c3:
            if risk_percentage > 30:
                st.metric("Commercial Health Score", f"{100 - risk_percentage}% (High Risk Profile)", delta="-Alert Matrix", delta_color="inverse")
            else:
                st.metric("Commercial Health Score", f"{100 - risk_percentage}% (Healthy Bid)", delta="Go Approved", delta_color="normal")
                
        st.write("---")
        st.markdown("💬 **Interactive Workspace:** Edit fields directly inside the cells, modify assigned teams, or flip the workflow status on your mobile screen.")
        
        # Stable, non-recursive data rendering layer
        st.data_editor(
            current_df,
            column_config={
                "category": st.column_config.SelectboxColumn(
                    "Structural Workspace",
                    options=["Proposal Draft", "Risk Audit", "Calculation Layer"],
                    required=True,
                ),
                "assigned_team": st.column_config.SelectboxColumn(
                    "Assigned Department",
                    options=["Sales Planning", "Legal/Compliance", "Technical Ops", "Commercial Finance"],
                    required=True,
                ),
                "status": st.column_config.SelectboxColumn(
                    "Review Status Workflow",
                    options=["Pending Review", "Approved", "Escalated/Risk Flagged"],
                    required=True,
                ),
                "risk_multiplier": st.column_config.NumberColumn(
                    "Risk Factor Score",
                    min_value=1.00,
                    max_value=3.00,
                    step=0.05,
                    format="%.2f",
                ),
                "requirement_extracted": st.column_config.TextColumn(
                    "Extracted Core Requirements Text",
                    width="large"
                )
            },
            hide_index=True,
            use_container_width=True,
            on_change=sync_matrix_edits,       
            key="interactive_matrix_editor"     
        )
        
        st.write("---")
        st.subheader("📥 Export & Distribution Workspace")
        
        excel_payload = generate_enterprise_excel(current_df)
        
        st.download_button(
            label="📥 Download Enterprise Commercial Matrix (.xlsx Format)",
            data=excel_payload,
            file_name="finalized_commercial_rfp_matrix.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            type="primary"
        )
        
    else:
        st.info("Awaiting file upload. Drop your PDF, Word document, or Excel matrix on the left side to begin.")
